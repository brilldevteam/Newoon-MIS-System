from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "source" / "Blank Format.docx"
OUTPUT = ROOT / "kyc-part-1-template.docx"


def set_cell(table, row, col, text):
    table.rows[row].cells[col].text = text


def insert_paragraph_after(paragraph, text):
    new_paragraph = OxmlElement("w:p")
    paragraph._p.addnext(new_paragraph)
    wrapped = Paragraph(new_paragraph, paragraph._parent)
    wrapped.text = text
    return wrapped


def remove_rows_after(table, row_index):
    for row in list(table.rows[row_index + 1 :]):
        table._tbl.remove(row._tr)


def main():
    document = Document(SOURCE)

    document.paragraphs[0].text = "Date: {date}"
    document.paragraphs[1].text = "Reference: {reference}"

    general = document.tables[0]
    set_cell(general, 0, 2, "{legalName}")
    set_cell(general, 1, 2, "{commercialRegistrationNo}")
    set_cell(general, 2, 2, "{taxIdentificationNo}")
    set_cell(general, 3, 2, "Date: {dateOfIncorporation} | Country: {countryOfIncorporation}")
    set_cell(general, 4, 2, "{legalForm}")
    set_cell(general, 5, 2, "{registeredOfficeAddress}")
    set_cell(general, 6, 2, "{telephone}")
    set_cell(general, 7, 2, "{email}")
    set_cell(general, 8, 2, "{website}")
    set_cell(general, 9, 2, "{businessNature}")
    set_cell(general, 10, 2, "{licenseActivities}")
    set_cell(general, 11, 2, "{relatedIndustry}")
    set_cell(general, 12, 2, "{prospectiveService}")

    shareholders = document.tables[2]
    for row in range(1, 8):
        for col in range(7):
            set_cell(shareholders, row, col, "")
    shareholder_placeholders = [
        "{#shareholders}{shareholderNo}",
        "{shareholderFullName}",
        "{shareholderNationality}",
        "{shareholderDateOfBirth}",
        "{shareholderIdentityNumber}",
        "{shareholderOwnershipPercentage}",
        "{shareholderResidenceAddress}{/shareholders}",
    ]
    for col, placeholder in enumerate(shareholder_placeholders):
        set_cell(shareholders, 1, col, placeholder)
    set_cell(shareholders, 8, 6, "{totalOwnershipPercentage}")

    document.paragraphs[10].text = "Yes {uboDifferentYes}    No {uboDifferentNo}"
    document.paragraphs[11].text = "If yes, please provide UBOs/Group structure and details. UBOs/ Group structure: {uboGroupStructureNotes}"

    ubos = document.tables[3]
    ubo_placeholders = [
        "{#ubos}{uboNo}",
        "{uboFullName}",
        "{uboNationality}",
        "{uboDateOfBirth}",
        "{uboIdentityNumber}",
        "{uboOwnershipPercentage}",
        "{uboResidenceAddress}{/ubos}",
    ]
    for col, placeholder in enumerate(ubo_placeholders):
        set_cell(ubos, 1, col, placeholder)
    set_cell(ubos, 2, 6, "{totalUboPercentage}")

    managers = document.tables[4]
    for row in range(2, 30):
        for col in range(8):
            set_cell(managers, row, col, "")
    manager_placeholders = [
        "{#managers}{managerNo}",
        "{managerFullName}",
        "{managerEntityName}",
        "{managerNationalityAndAddress}",
        "{managerDateOfBirth}",
        "{managerIdentityNumber}",
        "{managerPosition}",
        "{managerAuthorizedSignatory}{/managers}",
    ]
    for col, placeholder in enumerate(manager_placeholders):
        set_cell(managers, 2, col, placeholder)
    remove_rows_after(managers, 2)

    document.paragraphs[31].text = "Yes, if yes, please provide full details: {pepQuestion}"
    document.paragraphs[32].text = "{pepDetails}"
    document.paragraphs[34].text = "No {pepNo}"
    document.paragraphs[37].text = "Yes, if yes, please provide full details of any sanctions applied and the current state of such sanctions: {sanctionQuestion}"
    document.paragraphs[39].text = "{sanctionDetails}"
    document.paragraphs[42].text = "No {sanctionNo}"
    document.paragraphs[47].text = "Yes, if yes, please provide details and passport copy: {dualCitizenshipQuestion}"
    document.paragraphs[48].text = "{dualCitizenshipDetails} | Passport copy: {dualCitizenshipPassportFileName}"
    document.paragraphs[51].text = "No {dualCitizenshipNo}"

    contact = document.tables[5]
    while len(contact.columns) < 4:
        contact.add_column(Inches(1.5))
    contact_widths = [Inches(1.45), Inches(1.75), Inches(1.45), Inches(1.75)]
    for row in contact.rows:
        for col_index, width in enumerate(contact_widths):
            row.cells[col_index].width = width
    contact_rows = [
        ("Full name", "{communicationFullName}", "Position / Job title", "{communicationPosition}"),
        ("Nationality", "{communicationNationality}", "QID / Passport Number", "{communicationIdentityNumber}"),
        ("Mobile Number", "{communicationMobile}", "Email", "{communicationEmail}"),
    ]
    for row_index, row_values in enumerate(contact_rows):
        for col_index, value in enumerate(row_values):
            set_cell(contact, row_index, col_index, value)
    for row_index in range(3, len(contact.rows)):
        for col_index in range(len(contact.columns)):
            set_cell(contact, row_index, col_index, "")

    checklist_indexes = list(range(57, 66))
    checklist_placeholders = [
        "{docCommercialRegistration}",
        "{docEntityCard}",
        "{docCertificateOfIncorporation}",
        "{docArticlesOfAssociation}",
        "{docIdentityCopies}",
        "{docLegalEntityShareholderCr}",
        "{docNationalAddressCertificates}",
        "{docAuditedFinancialStatements}",
        "{docTaxCard}",
    ]
    for paragraph_index, placeholder in zip(checklist_indexes, checklist_placeholders):
        document.paragraphs[paragraph_index].text = f"{placeholder} " + document.paragraphs[paragraph_index].text

    tax_card_paragraph = document.paragraphs[checklist_indexes[-1]]
    note_paragraph = insert_paragraph_after(tax_card_paragraph, "Additional documents: {additionalDocumentsText}")
    insert_paragraph_after(note_paragraph, "Additional notes for KYC preparation documents: {uploadedFilesNote}")

    declaration = document.tables[6]
    set_cell(declaration, 0, 1, "{declarationFullName}")
    set_cell(declaration, 1, 1, "{declarationPosition}")
    set_cell(declaration, 2, 1, "{declarationDate}")

    signature = document.tables[7]
    set_cell(signature, 0, 1, "{signatureFileName}")
    set_cell(signature, 1, 1, "{stampFileName}")

    aml = document.tables[8]
    set_cell(aml, 0, 1, "Accuracy checked: {amlAccuracyChecked} | Findings: {amlClarificationFindings} | Risk: {riskClassification} | Due diligence: {dueDiligenceType}")
    set_cell(aml, 1, 0, "AML Supervisor Name:")
    set_cell(aml, 1, 1, "{amlName}")
    set_cell(aml, 2, 0, "AML Supervisor signature:")
    set_cell(aml, 2, 1, "{amlSignatureFileName}")
    set_cell(aml, 3, 1, "{amlDate}")

    dmlro = document.tables[9]
    set_cell(dmlro, 1, 1, "{dmlroName}")
    set_cell(dmlro, 2, 1, "{dmlroSignatureFileName}")
    set_cell(dmlro, 3, 1, "{dmlroDate}")
    set_cell(dmlro, 4, 0, "Decision:")
    set_cell(dmlro, 4, 1, "{dmlroDecision}")
    for label, placeholder in [
        ("Conditions:", "{dmlroConditions}"),
        ("Reason / additional information:", "{dmlroReason}"),
        ("Comments:", "{dmlroComments}"),
    ]:
        row = dmlro.add_row()
        row.cells[0].text = label
        row.cells[1].text = placeholder

    mlro = document.tables[10]
    set_cell(mlro, 1, 1, "{mlroName}")
    set_cell(mlro, 2, 1, "{mlroSignatureFileName}")
    set_cell(mlro, 3, 1, "{mlroDate}")
    set_cell(mlro, 4, 0, "Final decision:")
    set_cell(mlro, 4, 1, "{mlroDecision}")
    for label, placeholder in [
        ("Final risk classification:", "{mlroFinalRiskClassification}"),
        ("Risk reason category:", "{mlroRiskReasonCategory}"),
        ("Risk explanation:", "{mlroRiskExplanation}"),
        ("Conditions:", "{mlroConditions}"),
        ("Comments:", "{mlroComments}"),
    ]:
        row = mlro.add_row()
        row.cells[0].text = label
        row.cells[1].text = placeholder

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
